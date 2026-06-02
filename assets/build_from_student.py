# -*- coding: utf-8 -*-
"""Build the final DOCX report using the student template as the base.

Strategy:
1. Start with student_template.docx which has all the GOST elements
   (header1.xml with frame+small stamp drawing, body anchor with the
   large stamp on the contents page, multi-section layout).
2. Substitute the student-specific texts (name, group, theme, dates,
   supervisor) everywhere they appear: title page paragraphs,
   header1.xml, and the body anchor.
3. Clear all paragraphs of the body section EXCEPT the first one (which
   contains the large-stamp anchor) so the headers/frame/anchor stay
   intact.
4. Build new body content (СОДЕРЖАНИЕ, ВВЕДЕНИЕ, sections, conclusion,
   bibliography) and insert it after the anchor paragraph.
"""

import copy
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import content as C


WORK_DIR = Path('/home/ubuntu/report_work')
SRC_TEMPLATE = WORK_DIR / 'student_template.docx'
FIG_DIR = WORK_DIR / 'figures'
SCREEN_DIR = WORK_DIR / 'screens'
OUT_PATH = WORK_DIR / 'Балыкбаев_отчет_BI_визуализация.docx'

# Text substitutions - applied to document.xml.
# Maps original student-template strings to ours. Order matters: longer
# strings come first so they win.
SUBSTITUTIONS = [
    # Theme / название практики
    ('«QI Инжиниринг»', '«Базы данных»'),
    ('QI Инжиниринг', 'Базы данных'),
    # Student full name (title page)
    ('Азатов Радмир Кахарманович', 'Балыкбаев Бауыржан Шоканулы'),
    # Short name in stamps (Выполнил)
    ('Азатов Р.К.', 'Балыкбаев Б.Ш.'),
    # Supervisor full name (title page)
    ('Аубакиров Еркебулан Еркинович', 'Бычкова Светлана Сергеевна'),
    # Supervisor short name in stamps
    ('Аубакиров Е.Е.', 'Бычкова С.С.'),
    # Group
    ('Группа: РПО3 21-Р', 'Группа: РПО7-24Р'),
    ('РПО3 21-Р', 'РПО7-24Р'),
    # Course
    ('Курс: 4', 'Курс: 2'),
    # Period dates (title page)
    ('«09» октября 2024', '«06» мая 2026'),
    ('«24» октября 2024', '«19» мая 2026'),
    # Литера "У" + Листов "25" → 34. The "У25" pair is uniquely from
    # the large-stamp anchor.
    ('У25', 'У34'),
]


def w(tag):
    return qn('w:' + tag)


def make_el(tag, **attrs):
    el = OxmlElement('w:' + tag)
    for k, v in attrs.items():
        el.set(qn('w:' + k), str(v))
    return el


# ---------------------------------------------------------------------------
# Run/paragraph helpers (copied from build_docx.py, simplified)
# ---------------------------------------------------------------------------

def _set_run_font(run, name='Times New Roman', size_pt=14, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(w('rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for a in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rFonts.set(qn('w:' + a), name)
    for tag in ('b', 'bCs'):
        for el in rPr.findall(w(tag)):
            rPr.remove(el)
    for tag in ('i', 'iCs'):
        for el in rPr.findall(w(tag)):
            rPr.remove(el)
    if italic:
        for tag in ('i', 'iCs'):
            el = OxmlElement('w:' + tag)
            rPr.append(el)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_body_paragraph(parent_el, text, *, first_line_cm=1.25,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       size_pt=14, italic=False, before_el=None):
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = make_el('spacing', after='0', before='0', line='240',
                      lineRule='auto')
    pPr.append(spacing)
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pPr.append(make_el('jc', val='both'))
    elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
        pPr.append(make_el('jc', val='center'))
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        pPr.append(make_el('jc', val='right'))
    if first_line_cm and first_line_cm > 0:
        twips = int(first_line_cm * 567)
        pPr.append(make_el('ind', firstLine=str(twips)))
    p_el.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = make_el('rFonts', ascii='Times New Roman', hAnsi='Times New Roman',
                     cs='Times New Roman', eastAsia='Times New Roman')
    rPr.append(rFonts)
    rPr.append(make_el('sz', val=str(size_pt * 2)))
    rPr.append(make_el('szCs', val=str(size_pt * 2)))
    if italic:
        rPr.append(OxmlElement('w:i'))
        rPr.append(OxmlElement('w:iCs'))
    color = make_el('color', val='000000')
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text or ''
    r.append(t)
    p_el.append(r)

    if before_el is not None:
        before_el.addprevious(p_el)
    else:
        parent_el.append(p_el)
    return p_el


def add_page_break_paragraph(parent_el, before_el=None):
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = make_el('spacing', after='0', before='0', line='240',
                      lineRule='auto')
    pPr.append(spacing)
    p_el.append(pPr)
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p_el.append(r)
    if before_el is not None:
        before_el.addprevious(p_el)
    else:
        parent_el.append(p_el)
    return p_el


def add_image_paragraph(parent_el, doc, image_path, *, width_cm=15.5,
                        before_el=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    _set_run_font(run, size_pt=14)

    p_el = p._p
    p_el.getparent().remove(p_el)
    if before_el is not None:
        before_el.addprevious(p_el)
    else:
        parent_el.append(p_el)
    return p_el


# ---------------------------------------------------------------------------
# Text substitution helpers
# ---------------------------------------------------------------------------

def _do_substitution_in_runs(paragraph_xml):
    """Apply SUBSTITUTIONS across all <w:t> elements of a paragraph XML node.

    First concatenates text of all runs, applies substitutions, then writes
    the result back. We can lose run-level formatting (we put all the
    substituted text into the first run). For title/stamps where the styles
    follow paragraph-level rPr this is acceptable.
    """
    # Find all <w:t> children
    ts = paragraph_xml.findall('.//' + w('t'))
    if not ts:
        return False
    full = ''.join(t.text or '' for t in ts)
    new = full
    for src, dst in SUBSTITUTIONS:
        new = new.replace(src, dst)
    if new == full:
        return False
    # Put new text into first <w:t>, blank the rest
    ts[0].text = new
    ts[0].set(qn('xml:space'), 'preserve')
    for t in ts[1:]:
        t.text = ''
    return True


def substitute_text_in_doc(doc):
    """Walk every paragraph in the document and substitute texts."""
    body = doc.element.body
    # Substitute in all paragraphs (top-level and inside tables/anchors)
    n = 0
    for p_el in body.iter(w('p')):
        if _do_substitution_in_runs(p_el):
            n += 1
    return n


def substitute_text_in_xml_file(xml_path):
    """Apply substitutions to an XML file at character level.

    Used for header1.xml, where text may be split across many runs
    inside a complex grouped shape.
    """
    txt = xml_path.read_text(encoding='utf-8')
    original = txt
    # We work at the run/<w:t> level rather than raw string: parse the
    # text by concatenating siblings within each paragraph and grouping
    # them, similar to substitute_text_in_doc. For simplicity we do it
    # with regex - find each paragraph, then substitute texts inside.
    def replace_in_para(m):
        para = m.group(0)
        # Extract all <w:t...>...</w:t> contents in order
        t_pattern = re.compile(r'(<w:t[^>]*>)([^<]*)(</w:t>)')
        ts = list(t_pattern.finditer(para))
        if not ts:
            return para
        full = ''.join(t.group(2) for t in ts)
        new = full
        for src, dst in SUBSTITUTIONS:
            new = new.replace(src, dst)
        if new == full:
            return para
        # Place the new full text into the FIRST w:t, blank the rest.
        first = ts[0]
        new_para = (para[:first.start()] +
                    first.group(1) + new + first.group(3))
        # Now keep everything after first and blank the inner texts.
        rest_start = first.end()
        rest = para[rest_start:]
        # Replace each subsequent <w:t>...</w:t> with empty content.
        rest = re.sub(r'(<w:t[^>]*>)([^<]*)(</w:t>)', r'\1\3', rest)
        return new_para + rest

    new_txt = re.sub(r'<w:p\b.*?</w:p>', replace_in_para, txt, flags=re.DOTALL)
    if new_txt != original:
        xml_path.write_text(new_txt, encoding='utf-8')
        return True
    return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    shutil.copy(SRC_TEMPLATE, OUT_PATH)
    doc = Document(str(OUT_PATH))

    body = doc.element.body
    # All sectPr blocks (some are inside paragraphs for continuous sections)
    sect_elems = list(body.iter(w('sectPr')))
    print(f'Found {len(sect_elems)} sectPr blocks')
    # The final sectPr is the direct child of body
    final_sect = None
    for c in body:
        if c.tag == w('sectPr'):
            final_sect = c
    assert final_sect is not None, 'No body-level sectPr found'

    # Substitute texts in the in-memory doc (covers title page,
    # tables, and the body anchor).
    n = substitute_text_in_doc(doc)
    print(f'Substituted text in {n} paragraphs (body)')

    # The first paragraph in section 2 (body) holds the LARGE-stamp anchor.
    # Find it: after the second sectPr, the first <w:p> is the anchor para.
    # Section 2 in python-docx is doc.sections[2]; its content area starts
    # right after the previous sectPr.
    # We need to clear EVERYTHING in section 2 except: the anchor paragraph
    # (first one), final sectPr, and we'll add our own content.

    # Iterate body children; find the index of the FIRST paragraph in
    # section 2 (= immediately after second sectPr).
    # The second sectPr lives inside a paragraph (sectPr inside pPr), so
    # actually section 2 begins after the paragraph that contains sectPr.
    # python-docx pattern: sectPr is at the end of last paragraph of a
    # section; for "continuous" type it lives inside a paragraph too.
    # Easier: find the paragraph that contains the LARGE-stamp anchor
    # by searching for <wp:anchor>.
    anchor_para = None
    for p_el in body.iter(w('p')):
        if p_el.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor') is not None:
            anchor_para = p_el
            break
    assert anchor_para is not None, 'Failed to find body anchor paragraph'
    print(f'Anchor paragraph located')

    # Remove every body child that comes AFTER anchor_para except the
    # final sectPr. (The final sectPr is a child of <w:body>, not of any
    # paragraph.)
    anchor_idx = list(body).index(anchor_para)
    final_sect_idx = list(body).index(final_sect)
    print(f'Anchor at index {anchor_idx}, final sectPr at {final_sect_idx}')
    to_remove = list(body)[anchor_idx + 1:final_sect_idx]
    print(f'Removing {len(to_remove)} elements between anchor and final sectPr')
    for el in to_remove:
        body.remove(el)

    # Now anchor_para is followed directly by final_sect.
    # We'll insert content before final_sect (final_sect.addprevious).
    SECT = final_sect

    # Ensure final section's page margins are correct.
    pgMar = final_sect.find(w('pgMar'))
    if pgMar is not None:
        pgMar.set(qn('w:left'), str(int(30 * 56.7)))
        pgMar.set(qn('w:right'), str(int(10 * 56.7)))
        pgMar.set(qn('w:top'), str(int(15 * 56.7)))
        pgMar.set(qn('w:bottom'), str(int(30 * 56.7)))
        pgMar.set(qn('w:header'), '720')
        pgMar.set(qn('w:footer'), '720')

    # The anchor paragraph itself may have leftover w:r runs with style
    # info (TOC heading style "1"). Strip those runs but keep the anchor
    # drawing run. Actually the anchor para in the student template has
    # ONLY a single run that wraps the drawing. We'll add a w:r with the
    # СОДЕРЖАНИЕ heading text after the drawing run so the page starts
    # with the heading on top.
    # Simpler: insert a centered heading paragraph BEFORE inserting other
    # contents. But that paragraph would appear before the anchor, which
    # is fine since the anchor is positioned absolutely.

    # Counters
    fig_counters = {'1': 0, '2': 0}
    table_counters = {'1': 0, '2': 0}

    state = {'first_struct_heading': True}

    def heading_struct(text):
        if state['first_struct_heading']:
            state['first_struct_heading'] = False
        else:
            add_page_break_paragraph(body, before_el=SECT)
        add_body_paragraph(body, text, first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           before_el=SECT)
        add_body_paragraph(body, '', first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)

    def heading_section(text):
        add_page_break_paragraph(body, before_el=SECT)
        add_body_paragraph(body, text, first_line_cm=1.25,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)
        add_body_paragraph(body, '', first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)

    def heading_subsection(text):
        add_body_paragraph(body, '', first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)
        add_body_paragraph(body, text, first_line_cm=1.25,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)
        add_body_paragraph(body, '', first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)

    def add_figure(img_relpath, caption, sect_num):
        fig_counters[sect_num] += 1
        n = f"{sect_num}.{fig_counters[sect_num]}"
        if img_relpath.startswith('screens/'):
            path = SCREEN_DIR / img_relpath.split('/', 1)[1]
            width_cm = 15.0
        else:
            path = FIG_DIR / img_relpath
            width_cm = 14.0
        add_image_paragraph(body, doc, path, width_cm=width_cm,
                            before_el=SECT)
        cap_text = f"Рисунок-{n} {caption}"
        add_body_paragraph(body, cap_text, first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           size_pt=12, italic=True, before_el=SECT)
        add_body_paragraph(body, '', first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)

    def add_table(caption, rows, sect_num):
        table_counters[sect_num] += 1
        n = f"{sect_num}.{table_counters[sect_num]}"
        cap_text = f"Таблица-{n} {caption}"
        add_body_paragraph(body, cap_text, first_line_cm=1.25,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)
        n_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.alignment = 1
        tbl = table._tbl
        tblPr = tbl.find(w('tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for tag in ('tblBorders',):
            for el in tblPr.findall(w(tag)):
                tblPr.remove(el)
        borders = OxmlElement('w:tblBorders')
        for b in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elb = OxmlElement('w:' + b)
            elb.set(qn('w:val'), 'single')
            elb.set(qn('w:sz'), '4')
            elb.set(qn('w:color'), '000000')
            borders.append(elb)
        tblPr.append(borders)
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i == 0
                               else WD_ALIGN_PARAGRAPH.LEFT)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(cell_text)
                _set_run_font(run, size_pt=12)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        tbl_el = table._tbl
        tbl_el.getparent().remove(tbl_el)
        SECT.addprevious(tbl_el)
        add_body_paragraph(body, '', first_line_cm=0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           before_el=SECT)

    def render_items(items, sect_num):
        for item in items:
            kind = item[0]
            if kind == 'p':
                add_body_paragraph(body, item[1], first_line_cm=1.25,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   before_el=SECT)
            elif kind == 'fig':
                add_figure(item[1], item[2], sect_num)
            elif kind == 'table_h':
                add_table(item[1], item[2], sect_num)
            elif kind == 'sp':
                add_body_paragraph(body, '', first_line_cm=0,
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                   before_el=SECT)
            else:
                raise ValueError(f'Unknown item kind: {kind}')

    # ------------------- СОДЕРЖАНИЕ -------------------
    heading_struct('СОДЕРЖАНИЕ')
    contents_entries = [
        ('ВВЕДЕНИЕ', 3),
        ('1 ОСНОВНАЯ ЧАСТЬ', 4),
        ('1.1 Основные концепции BI и визуализации данных', 4),
        ('1.2 Power BI: архитектура и возможности', 7),
        ('1.3 Tableau: функциональность и применение', 10),
        ('1.4 Metabase: открытый инструмент BI', 13),
        ('1.5 Сравнение и выбор инструмента', 18),
        ('2 СПЕЦИАЛЬНАЯ ЧАСТЬ', 20),
        ('2.1 Создание ETL-процессов для миграции данных', 20),
        ('2.2 Облачные хранилища данных', 23),
        ('2.3 Работа с большими данными. Hadoop и Spark', 25),
        ('2.4 Оптимизация запросов в больших проектах', 27),
        ('ЗАКЛЮЧЕНИЕ', 30),
        ('СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ', 31),
    ]
    for title, page in contents_entries:
        p_el = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pPr.append(make_el('spacing', after='0', before='0', line='240',
                           lineRule='auto'))
        tabs = OxmlElement('w:tabs')
        tab = make_el('tab', val='right', leader='dot', pos='9070')
        tabs.append(tab)
        pPr.append(tabs)
        p_el.append(pPr)

        def _run(text):
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rPr.append(make_el('rFonts', ascii='Times New Roman',
                               hAnsi='Times New Roman',
                               cs='Times New Roman',
                               eastAsia='Times New Roman'))
            rPr.append(make_el('sz', val='28'))
            rPr.append(make_el('szCs', val='28'))
            rPr.append(make_el('color', val='000000'))
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            r.append(t)
            return r

        p_el.append(_run(title))
        r_tab = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(make_el('rFonts', ascii='Times New Roman',
                           hAnsi='Times New Roman',
                           cs='Times New Roman'))
        rPr.append(make_el('sz', val='28'))
        rPr.append(make_el('szCs', val='28'))
        r_tab.append(rPr)
        r_tab.append(OxmlElement('w:tab'))
        p_el.append(r_tab)
        p_el.append(_run(str(page)))

        SECT.addprevious(p_el)

    # ------------------- ВВЕДЕНИЕ ----------------------
    heading_struct('ВВЕДЕНИЕ')
    render_items(C.INTRO, '1')

    # ------------------- 1 ОСНОВНАЯ ЧАСТЬ --------------
    heading_section('1 ОСНОВНАЯ ЧАСТЬ')
    heading_subsection('1.1 Основные концепции BI и визуализации данных')
    render_items(C.S_1_1, '1')

    heading_subsection('1.2 Power BI: архитектура и возможности')
    render_items(C.S_1_2, '1')

    heading_subsection('1.3 Tableau: функциональность и применение')
    render_items(C.S_1_3, '1')

    heading_subsection('1.4 Metabase: открытый инструмент BI')
    render_items(C.S_1_4, '1')

    heading_subsection('1.5 Сравнение и выбор инструмента')
    render_items(C.S_1_5, '1')

    # ------------------- 2 СПЕЦИАЛЬНАЯ ЧАСТЬ -----------
    heading_section('2 СПЕЦИАЛЬНАЯ ЧАСТЬ')
    heading_subsection('2.1 Создание ETL-процессов для миграции данных')
    render_items(C.S_2_1, '2')

    heading_subsection('2.2 Облачные хранилища данных')
    render_items(C.S_2_2, '2')

    heading_subsection('2.3 Работа с большими данными. Hadoop и Spark')
    render_items(C.S_2_3, '2')

    heading_subsection('2.4 Оптимизация запросов в больших проектах')
    render_items(C.S_2_4, '2')

    # ------------------- ЗАКЛЮЧЕНИЕ --------------------
    heading_struct('ЗАКЛЮЧЕНИЕ')
    render_items(C.CONCLUSION, '2')

    # ------------------- СПИСОК ЛИТЕРАТУРЫ -------------
    heading_struct('СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ')
    for i, src in enumerate(C.BIBLIO, 1):
        text = f"{i}. {src}"
        add_body_paragraph(body, text, first_line_cm=1.25,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           before_el=SECT)

    doc.save(str(OUT_PATH))
    print(f'Saved (pre-header-edit): {OUT_PATH}')

    # ------------------------------------------------------------------
    # Post-process: apply same text substitutions to header1.xml because
    # python-docx does not expose header XML through Document objects
    # consistently for complex grouped shapes.
    # ------------------------------------------------------------------
    tmp = WORK_DIR / '_unzip_tmp'
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(OUT_PATH, 'r') as z:
        z.extractall(tmp)
    header1_path = tmp / 'word' / 'header1.xml'
    if header1_path.exists():
        if substitute_text_in_xml_file(header1_path):
            print('Substituted text in header1.xml')
    # Rezip
    new_out = OUT_PATH.with_suffix('.docx.new')
    with zipfile.ZipFile(new_out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for f in tmp.rglob('*'):
            if f.is_file():
                zout.write(f, f.relative_to(tmp))
    shutil.move(new_out, OUT_PATH)
    shutil.rmtree(tmp)
    print(f'Final saved: {OUT_PATH}')


if __name__ == '__main__':
    main()
